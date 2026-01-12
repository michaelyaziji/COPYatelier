"""Email service for sending documents via SMTP."""

import io
import logging
import smtplib
import ssl
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from typing import Optional

from .config import get_settings

logger = logging.getLogger(__name__)


def extract_clean_content(content: str) -> str:
    """
    Extract clean content from potentially JSON-wrapped output.

    Args:
        content: Raw content that may be wrapped in JSON

    Returns:
        Clean text content
    """
    import json
    import re

    cleaned = content.strip()

    # Remove markdown code fences (```json ... ``` or ``` ... ```)
    cleaned = re.sub(r'^```(?:json)?\s*\n?', '', cleaned)
    cleaned = re.sub(r'\n?```\s*$', '', cleaned)
    cleaned = cleaned.strip()

    # Try to parse as JSON and extract the "output" field
    try:
        if cleaned.startswith('{'):
            parsed = json.loads(cleaned)
            if 'output' in parsed:
                return parsed['output'].strip()
    except (json.JSONDecodeError, KeyError):
        pass

    # Check for "output": pattern even if not valid JSON
    match = re.search(r'"output"\s*:\s*"((?:[^"\\]|\\.)*)"', cleaned, re.DOTALL)
    if match:
        # Unescape JSON string
        result = match.group(1)
        result = result.replace('\\n', '\n')
        result = result.replace('\\t', '\t')
        result = result.replace('\\"', '"')
        result = result.replace('\\\\', '\\')
        return result.strip()

    # More aggressive extraction if "output" is found
    if '"output"' in cleaned:
        start = cleaned.find('"output"')
        if start != -1:
            after = cleaned[start:]
            colon_quote = after.find('": "')
            if colon_quote != -1:
                content_start = colon_quote + 4
                # Find closing quote, handling escapes
                i = content_start
                escaped = False
                while i < len(after):
                    if escaped:
                        escaped = False
                        i += 1
                        continue
                    if after[i] == '\\':
                        escaped = True
                        i += 1
                        continue
                    if after[i] == '"':
                        break
                    i += 1

                if i > content_start:
                    extracted = after[content_start:i]
                    extracted = extracted.replace('\\n', '\n')
                    extracted = extracted.replace('\\t', '\t')
                    extracted = extracted.replace('\\"', '"')
                    extracted = extracted.replace('\\\\', '\\')
                    return extracted.strip()

    return cleaned


def generate_word_document(content: str, title: Optional[str] = None) -> bytes:
    """
    Generate a professionally formatted Word document from text content.

    Args:
        content: The document content
        title: Optional document title

    Returns:
        Word document as bytes
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import re

    # Extract clean content first
    clean_content = extract_clean_content(content)

    doc = Document()

    # Set up document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Add title if provided
    if title:
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add a subtle line after title
        doc.add_paragraph()

    # Parse content into paragraphs
    lines = clean_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        trimmed = line.strip()

        # Skip empty lines (but track for paragraph spacing)
        if not trimmed:
            i += 1
            continue

        # Handle markdown-style headers
        if trimmed.startswith('### '):
            p = doc.add_heading(trimmed[4:], level=3)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif trimmed.startswith('## '):
            p = doc.add_heading(trimmed[3:], level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
        elif trimmed.startswith('# '):
            p = doc.add_heading(trimmed[2:], level=1)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(10)
        elif trimmed.startswith('- ') or trimmed.startswith('* '):
            # Bullet point
            p = doc.add_paragraph(trimmed[2:], style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
        elif re.match(r'^\d+\.\s', trimmed):
            # Numbered list
            text = re.sub(r'^\d+\.\s*', '', trimmed)
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.space_after = Pt(4)
        else:
            # Regular paragraph - handle inline formatting
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15

            # Process inline bold and italic markers
            add_formatted_text(p, trimmed)

        i += 1

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def add_formatted_text(paragraph, text: str):
    """
    Add text to a paragraph with markdown-style bold and italic formatting.

    Args:
        paragraph: The docx paragraph to add text to
        text: Text potentially containing **bold** and *italic* markers
    """
    from docx.shared import Pt
    import re

    # Pattern to match **bold**, *italic*, or ***bold italic***
    pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*[^*]+?\*)'

    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith('***') and part.endswith('***'):
            # Bold italic
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
            run.font.size = Pt(11)
        elif part.startswith('**') and part.endswith('**'):
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(11)
        else:
            # Regular text
            run = paragraph.add_run(part)
            run.font.size = Pt(11)


async def send_document_email(
    to_email: str,
    document_content: str,
    session_title: Optional[str] = None,
    personal_message: Optional[str] = None,
) -> bool:
    """
    Send a document via email using SMTP with Word attachment.

    Args:
        to_email: Recipient email address
        document_content: The document content to send
        session_title: Optional session title for the email subject
        personal_message: Optional personal message to include in the email

    Returns:
        True if email was sent successfully, False otherwise
    """
    settings = get_settings()

    if not settings.smtp_username or not settings.smtp_password:
        logger.warning("SMTP credentials not configured - email disabled")
        raise ValueError("Email service is not configured. Please set SMTP_USERNAME and SMTP_PASSWORD.")

    try:
        subject = f"Your Atelier Document: {session_title}" if session_title else "Your Atelier Document"

        # Generate Word document with title
        word_bytes = generate_word_document(document_content, title=session_title)

        # Create filename from title - clean and format nicely
        filename = "Atelier_Document"
        if session_title:
            # Clean the title for use as filename
            clean_title = ''.join(c if c.isalnum() or c in ' -_' else '' for c in session_title)
            clean_title = clean_title.strip().replace(' ', '_')[:50]
            if clean_title:
                filename = clean_title

        # Format the document nicely for email (preview in body)
        clean_preview = extract_clean_content(document_content)
        preview_content = clean_preview[:1000] + ('...' if len(clean_preview) > 1000 else '')

        # Create message container
        msg = MIMEMultipart('alternative')
        msg['Subject'] = subject
        msg['From'] = settings.smtp_from_email
        msg['To'] = to_email

        # Build personal message section if provided
        message_section_text = ""
        if personal_message:
            message_section_text = f"""{'-' * 40}
MESSAGE FROM SENDER:
{'-' * 40}

{personal_message}

"""

        # Plain text version
        text_content = f"""Your Document is Ready
{'=' * 40}

{session_title or 'Generated with Atelier AI Writing Studio'}

Word Document Attached - Open the attached .docx file for the full formatted document.

{message_section_text}{'-' * 40}
PREVIEW:
{'-' * 40}

{preview_content}

{'-' * 40}

✨ Made with Atelier — where AI writers and editors collaborate until it's right.
   https://atelierwritereditor.com
"""

        # Build personal message HTML section if provided
        message_section_html = ""
        if personal_message:
            # Escape HTML in the message and convert newlines to <br>
            escaped_message = (personal_message
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&#39;')
                .replace('\n', '<br>'))
            message_section_html = f"""
            <div class="personal-message">
                <div class="message-label">Message from sender</div>
                <div class="message-content">{escaped_message}</div>
            </div>
            """

        # HTML version
        html_content = f"""
        <html>
        <head>
            <style>
                body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; line-height: 1.6; color: #333; padding: 20px; }}
                .header {{ background: linear-gradient(135deg, #8b5cf6, #7c3aed); color: white; padding: 30px; border-radius: 12px; margin-bottom: 30px; }}
                .header h1 {{ margin: 0; font-size: 24px; }}
                .header p {{ margin: 10px 0 0 0; opacity: 0.9; }}
                .attachment-note {{ background: #ecfdf5; border: 1px solid #a7f3d0; border-radius: 8px; padding: 16px; margin-bottom: 20px; color: #065f46; }}
                .attachment-note strong {{ color: #047857; }}
                .personal-message {{ background: #fef3c7; border: 1px solid #fcd34d; border-radius: 8px; padding: 16px; margin-bottom: 20px; }}
                .message-label {{ font-size: 12px; color: #92400e; margin-bottom: 8px; text-transform: uppercase; letter-spacing: 0.5px; font-weight: 600; }}
                .message-content {{ color: #78350f; }}
                .preview {{ background: #f9fafb; border: 1px solid #e5e7eb; border-radius: 12px; padding: 30px; white-space: pre-wrap; font-family: inherit; }}
                .preview-label {{ font-size: 12px; color: #6b7280; margin-bottom: 10px; text-transform: uppercase; letter-spacing: 0.5px; }}
                .footer {{ margin-top: 30px; padding-top: 20px; border-top: 1px solid #e5e7eb; color: #6b7280; font-size: 14px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>Your Document is Ready</h1>
                <p>{session_title or 'Generated with Atelier AI Writing Studio'}</p>
            </div>
            <div class="attachment-note">
                <strong>Word Document Attached</strong> - Open the attached .docx file for the full formatted document.
            </div>
            {message_section_html}
            <div class="preview-label">Preview</div>
            <div class="preview">{preview_content}</div>
            <div class="footer">
                <p>✨ Made with <strong>Atelier</strong> — where AI writers and editors collaborate until it's right.</p>
                <p><a href="https://atelierwritereditor.com">Try Atelier</a></p>
            </div>
        </body>
        </html>
        """

        # Attach text and HTML parts
        part1 = MIMEText(text_content, 'plain')
        part2 = MIMEText(html_content, 'html')
        msg.attach(part1)
        msg.attach(part2)

        # Attach Word document
        attachment = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
        attachment.set_payload(word_bytes)
        encoders.encode_base64(attachment)
        attachment.add_header(
            'Content-Disposition',
            f'attachment; filename="{filename}.docx"'
        )
        msg.attach(attachment)

        # Send email via SMTP
        if settings.smtp_use_ssl:
            # Use SSL (port 465)
            context = ssl.create_default_context()
            with smtplib.SMTP_SSL(settings.smtp_host, settings.smtp_port, context=context) as server:
                server.login(settings.smtp_username, settings.smtp_password)
                server.sendmail(settings.smtp_from_email, to_email, msg.as_string())
        else:
            # Use TLS (port 587)
            with smtplib.SMTP(settings.smtp_host, settings.smtp_port) as server:
                server.starttls()
                server.login(settings.smtp_username, settings.smtp_password)
                server.sendmail(settings.smtp_from_email, to_email, msg.as_string())

        logger.info(f"Email sent successfully to {to_email} with attachment")
        return True

    except smtplib.SMTPAuthenticationError as e:
        logger.error(f"SMTP authentication failed: {e}")
        raise ValueError("Email authentication failed. Please check SMTP credentials.")
    except smtplib.SMTPException as e:
        logger.error(f"SMTP error sending email to {to_email}: {e}")
        raise ValueError(f"Failed to send email: {str(e)}")
    except Exception as e:
        logger.error(f"Failed to send email to {to_email}: {e}")
        raise ValueError(f"Failed to send email: {str(e)}")
